import requests
from bs4 import BeautifulSoup
from docx import Document

# Paso 1: Obtener HTML
url = "https://dashboard-tickets-uvwhd2e9or9g6ngyp54kfb.streamlit.app/"
headers = {'User-Agent': 'Mozilla/5.0'}
response = requests.get(url, headers=headers)
soup = BeautifulSoup(response.text, "html.parser")

# Paso 2: Crear documento Word
doc = Document()
doc.add_heading('KPIs de Soporte Técnico - Zendesk', 0)

# Paso 3: Extraer contenido relevante (títulos y párrafos)
for section in soup.find_all(['h2', 'h3', 'p']):
    text = section.get_text(strip=True)
    if section.name == 'h2':
        doc.add_heading(text, level=1)
    elif section.name == 'h3':
        doc.add_heading(text, level=2)
    elif section.name == 'p':
        doc.add_paragraph(text)

# Paso 4: Guardar
doc.save("kpis_soporte_zendesk.docx")
